from __future__ import annotations

import logging
from typing import List

from api.core.rag.extractor.base_extractor import (
    BaseExtractor,
    Document,
    FileValidationError,
)
from api.core.rag.extractor.entities.extract_setting import ExtractSetting

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    def __init__(self, file_path: str):
        self._file_path = file_path

    def supported_formats(self) -> list[str]:
        return [".docx"]

    def extract(self, extract_setting: ExtractSetting | None = None) -> List[Document]:
        self.validate_file(self._file_path)

        try:
            import docx
        except ImportError:
            raise FileValidationError(
                "python-docx is required for Word extraction. Install it with: pip install python-docx"
            )

        try:
            document = docx.Document(self._file_path)
        except Exception as e:
            raise FileValidationError(f"Failed to open Word file: {e}")

        paragraphs = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        page_content = "\n".join(paragraphs)

        return [
            Document(
                page_content=page_content,
                metadata={
                    "source": self._file_path,
                    "paragraph_count": len(paragraphs),
                },
            )
        ]
