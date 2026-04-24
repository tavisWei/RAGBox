"""Document processing service for chunking and metadata extraction."""

import re
import os
from typing import List, Optional, Dict, Any, Tuple
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

# Optional PDF/HTML parsing
try:
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pypdfium2 as pdfium

    PYPDFIUM_AVAILABLE = True
except ImportError:
    PYPDFIUM_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentError(Exception):
    """Raised when document processing fails."""

    pass


class UnsupportedFormatError(DocumentError):
    """Raised when document format is not supported."""

    pass


class ChunkingStrategy(str, Enum):
    """Supported chunking strategies."""

    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"  # Split by sentences/paragraphs
    RECURSIVE = "recursive"  # Split by hierarchical delimiters


@dataclass
class DocumentChunk:
    """A chunk of a document."""

    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""

    title: Optional[str] = None
    author: Optional[str] = None
    created_at: Optional[str] = None
    file_type: Optional[str] = None
    file_path: Optional[str] = None
    extraction_method: Optional[str] = None
    page_count: Optional[int] = None
    paragraph_count: Optional[int] = None
    char_count: int = 0
    word_count: int = 0


@dataclass
class ProcessedDocument:
    """A fully processed document with chunks and metadata."""

    metadata: DocumentMetadata
    chunks: List[DocumentChunk]


@dataclass
class ProcessorConfig:
    """Configuration for document processing."""

    chunk_size: int = 500  # Characters per chunk
    chunk_overlap: int = 100  # Overlap between chunks
    strategy: ChunkingStrategy = ChunkingStrategy.FIXED_SIZE
    extract_metadata: bool = True
    preserve_headings: bool = True
    max_threads: int = 4


class DocumentProcessor:
    """Service for processing documents into chunks."""

    def __init__(self, config: Optional[ProcessorConfig] = None):
        """
        Initialize document processor.

        Args:
            config: Processing configuration
        """
        self.config = config or ProcessorConfig()

    async def process_file(self, file_path: str) -> ProcessedDocument:
        """
        Process a document file and return chunks.

        Args:
            file_path: Path to the document file

        Returns:
            ProcessedDocument with chunks and metadata
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        section_spans: List[Tuple[int, int, Dict[str, Any]]] = []
        extraction_method = None

        # Extract text based on file type
        if suffix == ".txt":
            text = await self._read_txt(path)
        elif suffix in [".md", ".markdown"]:
            text = await self._read_markdown(path)
        elif suffix == ".pdf":
            sections, extraction_method = await self._read_pdf_sections(path)
            text, section_spans = self._join_sections(sections)
        elif suffix in [".docx", ".doc"]:
            sections, extraction_method = await self._read_word_sections(path)
            text, section_spans = self._join_sections(sections)
        elif suffix in [".html", ".htm"]:
            text = await self._read_html(path)
        else:
            raise UnsupportedFormatError(f"Unsupported file format: {suffix}")

        # Extract metadata
        metadata = DocumentMetadata(
            title=path.stem,
            file_type=suffix,
            file_path=str(path),
            extraction_method=extraction_method,
            page_count=self._max_metadata_value(section_spans, "page"),
            paragraph_count=self._max_metadata_value(section_spans, "paragraph_index"),
            char_count=len(text),
            word_count=len(text.split()),
        )

        if self.config.extract_metadata:
            metadata = self._extract_metadata(text, metadata)

        # Chunk the document
        chunks = await self._chunk_text(text, metadata, section_spans)

        return ProcessedDocument(metadata=metadata, chunks=chunks)

    async def process_text(
        self,
        text: str,
        metadata: Optional[DocumentMetadata] = None,
    ) -> ProcessedDocument:
        """
        Process raw text and return chunks.

        Args:
            text: Raw text content
            metadata: Optional pre-extracted metadata

        Returns:
            ProcessedDocument with chunks and metadata
        """
        doc_metadata = metadata or DocumentMetadata(
            char_count=len(text),
            word_count=len(text.split()),
        )

        chunks = await self._chunk_text(text, doc_metadata)

        return ProcessedDocument(metadata=doc_metadata, chunks=chunks)

    async def _read_txt(self, path: Path) -> str:
        """Read plain text file."""
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(
            ThreadPoolExecutor(max_workers=1),
            lambda: path.read_text(encoding="utf-8", errors="replace"),
        )
        return text

    async def _read_markdown(self, path: Path) -> str:
        """Read markdown file."""
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(
            ThreadPoolExecutor(max_workers=1),
            lambda: path.read_text(encoding="utf-8", errors="replace"),
        )
        return text

    async def _read_pdf(self, path: Path) -> str:
        """Read PDF file and extract text."""
        sections, _ = await self._read_pdf_sections(path)
        text, _spans = self._join_sections(sections)
        return text

    async def _read_pdf_sections(self, path: Path) -> Tuple[List[Dict[str, Any]], str]:
        if not PYPDF_AVAILABLE and not PYPDFIUM_AVAILABLE:
            raise UnsupportedFormatError(
                "PDF support requires pypdf or pypdfium2. Run: pip install pypdf"
            )

        loop = asyncio.get_event_loop()

        def _extract_with_pypdfium():
            pdf = pdfium.PdfDocument(str(path))
            sections = []
            try:
                for page_index in range(len(pdf)):
                    page = pdf[page_index]
                    textpage = page.get_textpage()
                    text = textpage.get_text_bounded().strip()
                    textpage.close()
                    page.close()
                    if text:
                        sections.append(
                            {
                                "text": text,
                                "metadata": {
                                    "source": str(path),
                                    "file_type": ".pdf",
                                    "page": page_index + 1,
                                    "page_index": page_index,
                                    "total_pages": len(pdf),
                                    "block_type": "page",
                                    "extraction_method": "pypdfium2",
                                },
                            }
                        )
            finally:
                pdf.close()
            return sections

        def _extract_with_pypdf():
            reader = pypdf.PdfReader(str(path))
            sections = []
            total_pages = len(reader.pages)
            for page_index, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    sections.append(
                        {
                            "text": text.strip(),
                            "metadata": {
                                "source": str(path),
                                "file_type": ".pdf",
                                "page": page_index + 1,
                                "page_index": page_index,
                                "total_pages": total_pages,
                                "block_type": "page",
                                "extraction_method": "pypdf",
                            },
                        }
                    )
            return sections

        if PYPDFIUM_AVAILABLE:
            sections = await loop.run_in_executor(
                ThreadPoolExecutor(max_workers=1), _extract_with_pypdfium
            )
            if sections:
                return sections, "pypdfium2"
        sections = await loop.run_in_executor(
            ThreadPoolExecutor(max_workers=1), _extract_with_pypdf
        )
        if not sections:
            raise UnsupportedFormatError(
                "PDF text extraction returned no text. Scanned/image-only PDFs require OCR, which is not enabled in this lightweight parser."
            )
        return sections, "pypdf"

    async def _read_html(self, path: Path) -> str:
        """Read HTML file and extract text."""
        if not BS4_AVAILABLE:
            raise UnsupportedFormatError(
                "HTML support requires beautifulsoup4. Run: pip install beautifulsoup4"
            )

        loop = asyncio.get_event_loop()

        def _extract_html():
            soup = BeautifulSoup(
                path.read_text(encoding="utf-8", errors="replace"), "html.parser"
            )
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator="\n", strip=True)

        text = await loop.run_in_executor(
            ThreadPoolExecutor(max_workers=1), _extract_html
        )
        return text

    async def _read_word(self, path: Path) -> str:
        sections, _ = await self._read_word_sections(path)
        text, _spans = self._join_sections(sections)
        return text

    async def _read_word_sections(self, path: Path) -> Tuple[List[Dict[str, Any]], str]:
        if not DOCX_AVAILABLE:
            raise UnsupportedFormatError(
                "Word support requires python-docx. Run: pip install python-docx"
            )

        if path.suffix.lower() == ".doc":
            raise UnsupportedFormatError(
                "Legacy .doc is not supported directly. Please convert to .docx"
            )

        loop = asyncio.get_event_loop()

        def _extract_docx():
            document = docx.Document(str(path))
            sections = []
            paragraph_index = 0
            table_index = 0
            for child in document.element.body.iterchildren():
                if child.tag.endswith("}p"):
                    paragraph = Paragraph(child, document)
                    text = paragraph.text.strip()
                    if text:
                        paragraph_index += 1
                        sections.append(
                            {
                                "text": text,
                                "metadata": {
                                    "source": str(path),
                                    "file_type": ".docx",
                                    "paragraph_index": paragraph_index,
                                    "block_type": "paragraph",
                                    "style": paragraph.style.name
                                    if paragraph.style
                                    else None,
                                    "extraction_method": "python-docx",
                                },
                            }
                        )
                elif child.tag.endswith("}tbl"):
                    table = Table(child, document)
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            rows.append("\t".join(cells))
                    if rows:
                        table_index += 1
                        sections.append(
                            {
                                "text": "\n".join(rows),
                                "metadata": {
                                    "source": str(path),
                                    "file_type": ".docx",
                                    "table_index": table_index,
                                    "block_type": "table",
                                    "extraction_method": "python-docx",
                                },
                            }
                        )
            return sections

        sections = await loop.run_in_executor(
            ThreadPoolExecutor(max_workers=1), _extract_docx
        )
        return sections, "python-docx"

    def _join_sections(
        self, sections: List[Dict[str, Any]]
    ) -> Tuple[str, List[Tuple[int, int, Dict[str, Any]]]]:
        parts: List[str] = []
        spans: List[Tuple[int, int, Dict[str, Any]]] = []
        cursor = 0
        for section in sections:
            text = (section.get("text") or "").strip()
            if not text:
                continue
            if parts:
                parts.append("\n\n")
                cursor += 2
            start = cursor
            parts.append(text)
            cursor += len(text)
            spans.append((start, cursor, section.get("metadata", {})))
        return "".join(parts), spans

    def _max_metadata_value(
        self, spans: List[Tuple[int, int, Dict[str, Any]]], key: str
    ) -> Optional[int]:
        values = [meta.get(key) for _start, _end, meta in spans if meta.get(key)]
        return max(values) if values else None

    def _extract_metadata(
        self,
        text: str,
        metadata: DocumentMetadata,
    ) -> DocumentMetadata:
        """Extract metadata from document text."""
        lines = text.split("\n")

        # Try to find title from first heading
        for line in lines[:10]:
            line = line.strip()
            if line.startswith("# "):
                metadata.title = line[2:].strip()
                break

        # Try to find author (simple heuristics)
        author_patterns = [
            r"^author:\s*(.+)$",
            r"^by\s+(.+)$",
            r"^\*\*Author:\*\*\s*(.+)$",
        ]
        for pattern in author_patterns:
            match = re.search(pattern, text[:500], re.MULTILINE | re.IGNORECASE)
            if match:
                metadata.author = match.group(1).strip()
                break

        return metadata

    async def _chunk_text(
        self,
        text: str,
        metadata: DocumentMetadata,
        section_spans: Optional[List[Tuple[int, int, Dict[str, Any]]]] = None,
    ) -> List[DocumentChunk]:
        """
        Chunk text based on configured strategy.

        Args:
            text: Text to chunk
            metadata: Document metadata

        Returns:
            List of DocumentChunk objects
        """
        if self.config.strategy == ChunkingStrategy.FIXED_SIZE:
            return await self._chunk_fixed_size(text, metadata, section_spans)
        elif self.config.strategy == ChunkingStrategy.SEMANTIC:
            return await self._chunk_semantic(text, metadata, section_spans)
        elif self.config.strategy == ChunkingStrategy.RECURSIVE:
            return await self._chunk_recursive(text, metadata, section_spans)
        else:
            return await self._chunk_fixed_size(text, metadata, section_spans)

    def _build_chunk_metadata(
        self,
        metadata: DocumentMetadata,
        start: int,
        end: int,
        section_spans: Optional[List[Tuple[int, int, Dict[str, Any]]]] = None,
    ) -> Dict[str, Any]:
        chunk_metadata: Dict[str, Any] = {
            "title": metadata.title,
            "file_type": metadata.file_type,
            "source": metadata.file_path,
            "extraction_method": metadata.extraction_method,
        }
        if section_spans:
            matched = [
                meta
                for span_start, span_end, meta in section_spans
                if span_start < end and span_end > start
            ]
            pages = sorted({m.get("page") for m in matched if m.get("page")})
            paragraphs = sorted(
                {m.get("paragraph_index") for m in matched if m.get("paragraph_index")}
            )
            block_types = sorted(
                {m.get("block_type") for m in matched if m.get("block_type")}
            )
            if pages:
                chunk_metadata["pages"] = pages
                if len(pages) == 1:
                    chunk_metadata["page"] = pages[0]
            if paragraphs:
                chunk_metadata["paragraphs"] = paragraphs
            if block_types:
                chunk_metadata["block_types"] = block_types
        return {k: v for k, v in chunk_metadata.items() if v is not None}

    async def _chunk_fixed_size(
        self,
        text: str,
        metadata: DocumentMetadata,
        section_spans: Optional[List[Tuple[int, int, Dict[str, Any]]]] = None,
    ) -> List[DocumentChunk]:
        """Chunk by fixed character size with overlap."""
        chunks = []
        chunk_size = self.config.chunk_size
        overlap = self.config.chunk_overlap

        start = 0
        index = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at word boundary
            if end < len(text):
                boundary = text[start:end].rfind(" ")
                if boundary > chunk_size // 2:
                    end = start + boundary

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    DocumentChunk(
                        content=chunk_text,
                        chunk_index=index,
                        start_char=start,
                        end_char=end,
                        metadata={
                            **self._build_chunk_metadata(
                                metadata, start, end, section_spans
                            )
                        },
                    )
                )
                index += 1

            start = end - overlap
            if start < 0:
                start = 0

        return chunks

    async def _chunk_semantic(
        self,
        text: str,
        metadata: DocumentMetadata,
        section_spans: Optional[List[Tuple[int, int, Dict[str, Any]]]] = None,
    ) -> List[DocumentChunk]:
        """Chunk by sentences and paragraphs."""
        # Split by paragraph first
        paragraphs = re.split(r"\n\s*\n", text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks = []
        current_chunk = ""
        current_start = 0
        index = 0

        for para in paragraphs:
            para_start = text.find(para, current_start)

            if len(current_chunk) + len(para) + 2 <= self.config.chunk_size:
                if current_chunk:
                    current_chunk += "\n\n"
                current_chunk += para
            else:
                # Save current chunk if not empty
                if current_chunk.strip():
                    chunks.append(
                        DocumentChunk(
                            content=current_chunk.strip(),
                            chunk_index=index,
                            start_char=current_start,
                            end_char=current_start + len(current_chunk),
                            metadata=self._build_chunk_metadata(
                                metadata,
                                current_start,
                                current_start + len(current_chunk),
                                section_spans,
                            ),
                        )
                    )
                    index += 1

                # Start new chunk
                current_chunk = para
                current_start = para_start

        # Don't forget last chunk
        if current_chunk.strip():
            chunks.append(
                DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_index=index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                    metadata=self._build_chunk_metadata(
                        metadata,
                        current_start,
                        current_start + len(current_chunk),
                        section_spans,
                    ),
                )
            )

        return chunks

    async def _chunk_recursive(
        self,
        text: str,
        metadata: DocumentMetadata,
        section_spans: Optional[List[Tuple[int, int, Dict[str, Any]]]] = None,
    ) -> List[DocumentChunk]:
        """Chunk by hierarchical delimiters (markdown headers, code blocks, etc.)."""
        # Try splitting by markdown headers first
        header_pattern = r"(?=^#{1,6}\s)"

        # Split while keeping headers
        parts = re.split(header_pattern, text, flags=re.MULTILINE)
        parts = [p.strip() for p in parts if p.strip()]

        chunks = []
        index = 0

        for part in parts:
            if len(part) <= self.config.chunk_size:
                chunks.append(
                    DocumentChunk(
                        content=part,
                        chunk_index=index,
                        start_char=text.find(part),
                        end_char=text.find(part) + len(part),
                        metadata=self._build_chunk_metadata(
                            metadata,
                            text.find(part),
                            text.find(part) + len(part),
                            section_spans,
                        ),
                    )
                )
                index += 1
            else:
                # Split large parts by fixed size
                sub_chunks = await self._chunk_fixed_size(part, metadata, section_spans)
                for sc in sub_chunks:
                    sc.chunk_index = index
                    index += 1
                chunks.extend(sub_chunks)

        return chunks


# Factory function
def create_document_processor(
    chunk_size: int = 500,
    chunk_overlap: int = 100,
    strategy: str = "fixed_size",
    **kwargs,
) -> "DocumentProcessor":
    """
    Create a DocumentProcessor with configuration.

    Args:
        chunk_size: Characters per chunk
        chunk_overlap: Overlap between chunks
        strategy: "fixed_size", "semantic", or "recursive"
        **kwargs: Additional arguments

    Returns:
        Configured DocumentProcessor
    """
    config = ProcessorConfig(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        strategy=ChunkingStrategy(strategy),
        **kwargs,
    )
    return DocumentProcessor(config=config)
